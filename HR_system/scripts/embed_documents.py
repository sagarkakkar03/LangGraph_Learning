import sys, os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
"""
Embed HR documents into a FAISS vector store.

Reads .docx files from DocsHR/, matches each to its metadata in the
hr_documents table, chunks the text, and builds a FAISS index with
rich metadata attached to every chunk for filtered retrieval.

Also back-fills content_text and chunks into the Postgres tables.

Usage:
    python embed_documents.py
"""

import os
import logging
from pathlib import Path

from docx import Document as DocxDocument
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.core.config import (
    DOCS_DIR,
    FAISS_INDEX_DIR,
    EMBEDDING_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)
from app.db.database import init_db, SessionLocal, HRDocument, HRDocumentChunk

logging.basicConfig(level=logging.INFO, format="%(asctime)s  %(levelname)s  %(message)s")
logger = logging.getLogger("embed_documents")


def read_docx(path: str) -> str:
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _find_docx_file(file_path_from_db: str) -> str | None:
    """Resolve the file_path stored in DB to an actual path under DocsHR/."""
    candidate = os.path.join(DOCS_DIR, file_path_from_db)
    if os.path.isfile(candidate):
        return candidate

    basename = os.path.basename(file_path_from_db)
    for root, _, files in os.walk(DOCS_DIR):
        if basename in files:
            return os.path.join(root, basename)
    return None


def build_index():
    init_db()
    db = SessionLocal()

    embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    all_docs: list[Document] = []
    processed = 0
    skipped = 0

    try:
        hr_docs = db.query(HRDocument).filter(HRDocument.is_active.is_(True)).all()
        logger.info("Found %d active HR documents in DB.", len(hr_docs))

        for hr_doc in hr_docs:
            docx_path = _find_docx_file(hr_doc.file_path)
            if not docx_path:
                logger.warning("File not found for %s: %s", hr_doc.doc_code, hr_doc.file_path)
                skipped += 1
                continue

            try:
                full_text = read_docx(docx_path)
            except Exception as exc:
                logger.warning("Cannot read %s: %s", docx_path, exc)
                skipped += 1
                continue

            if not full_text.strip():
                skipped += 1
                continue

            hr_doc.content_text = full_text

            db.query(HRDocumentChunk).filter(
                HRDocumentChunk.document_id == hr_doc.id
            ).delete()

            chunks = splitter.split_text(full_text)
            metadata_base = {
                "doc_code": hr_doc.doc_code,
                "base_code": hr_doc.base_code,
                "title": hr_doc.title,
                "country_code": hr_doc.country_code,
                "country": hr_doc.country,
                "doc_type": hr_doc.doc_type,
                "category": hr_doc.category,
                "document_family": hr_doc.document_family,
                "role_family": hr_doc.role_family or "",
                "escalation_email": hr_doc.escalation_email,
                "escalation_department": hr_doc.escalation_department,
                "document_url": hr_doc.document_url or "",
                "related_docs": ", ".join(hr_doc.related_docs or []),
                "keywords": hr_doc.keywords or "",
            }

            for i, chunk_text in enumerate(chunks):
                db_chunk = HRDocumentChunk(
                    document_id=hr_doc.id,
                    chunk_index=i,
                    chunk_text=chunk_text,
                    token_count=len(chunk_text.split()),
                )
                db.add(db_chunk)

                lc_doc = Document(
                    page_content=chunk_text,
                    metadata={**metadata_base, "chunk_index": i},
                )
                all_docs.append(lc_doc)

            processed += 1
            if processed % 25 == 0:
                logger.info("Processed %d documents...", processed)

        db.commit()
        logger.info(
            "Done reading. Processed %d docs, skipped %d, total %d chunks.",
            processed, skipped, len(all_docs),
        )

        if not all_docs:
            logger.error("No chunks to embed. Aborting.")
            return

        logger.info("Building FAISS index with %d chunks...", len(all_docs))
        vectorstore = FAISS.from_documents(all_docs, embeddings)

        os.makedirs(FAISS_INDEX_DIR, exist_ok=True)
        vectorstore.save_local(FAISS_INDEX_DIR)
        logger.info("FAISS index saved to %s", FAISS_INDEX_DIR)

    except Exception:
        db.rollback()
        raise
    finally:
        db.close()


if __name__ == "__main__":
    build_index()
